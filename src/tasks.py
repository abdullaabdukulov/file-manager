import asyncio
import io
import logging
from uuid import UUID

from celery import Celery
from docx import Document
from PyPDF2 import PdfReader
from sqlalchemy import insert, select

from src.config import settings
from src.database import async_session, fetch_one
from src.files.constants import FileType
from src.files.models import File, FileMetadata
from src.files.s3 import download_from_s3

logger = logging.getLogger(__name__)

app = Celery("tasks", broker=settings.REDIS_URL, backend=settings.REDIS_URL)

app.conf.update(
    task_serializer="json",
    accept_content=["json"],
    result_serializer="json",
    timezone="UTC",
    enable_utc=True,
)


def run_async(coroutine):
    """Run an async coroutine in the current event loop synchronously."""
    loop = asyncio.get_event_loop()
    if loop.is_closed():
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
    try:
        return loop.run_until_complete(coroutine)
    finally:
        if loop.is_running():
            loop.close()


@app.task
def extract_metadata(file_id: str, s3_key: str, file_type: str) -> None:
    """Extract metadata from a file and save to database."""
    logger.info(
        f"Starting metadata extraction for file_id: {file_id}, s3_key: {s3_key}"
    )
    try:
        # Run async download_from_s3 in sync context
        file_content = run_async(download_from_s3(s3_key))
        metadata = {}

        if file_type == FileType.PDF.value:
            pdf = PdfReader(io.BytesIO(file_content))
            info = pdf.metadata or {}
            metadata = {
                "file_id": file_id,
                "page_count": len(pdf.pages),
                "title": info.get("/Title", ""),
                "author": info.get("/Author", ""),
                "creation_date": info.get("/CreationDate", ""),
                "creator": info.get("/Creator", ""),
            }
        elif file_type == FileType.DOCX.value:
            doc = Document(io.BytesIO(file_content))
            metadata = {
                "file_id": file_id,
                "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "table_count": len(doc.tables),
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "creation_date": str(doc.core_properties.created)
                if doc.core_properties.created
                else "",
            }

        async def save_metadata():
            async with async_session() as session:
                async with session.begin():
                    file_exists = await fetch_one(
                        select(File).where(File.id == UUID(file_id))
                    )
                    if file_exists:
                        await session.execute(insert(FileMetadata).values(**metadata))
                        logger.info(f"Metadata saved for file_id: {file_id}")
                    else:
                        logger.warning(
                            f"File not found for metadata extraction: {file_id}"
                        )

        # Run save_metadata synchronously in the Celery task
        run_async(save_metadata())
    except Exception as e:
        logger.error(
            f"Metadata extraction failed for file_id: {file_id}, error: {str(e)}"
        )
        raise
